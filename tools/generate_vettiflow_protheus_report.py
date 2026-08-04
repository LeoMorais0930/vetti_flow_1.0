from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "relatorio_vettiflow_protheus_fastapi.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 105, 120)
LIGHT_FILL = "E8EEF5"
SOFT_FILL = "F4F6F9"


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.xpath("./w:shd")
    if shd:
        shd = shd[0]
    else:
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill", fill)


def set_table_geometry(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = BLUE if level < 3 else DARK_BLUE


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_number(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        hdr[index].text = header
        set_cell_fill(hdr[index], LIGHT_FILL)
        for paragraph in hdr[index].paragraphs:
            paragraph.runs[0].bold = True
    for row_values in rows:
        row = table.add_row().cells
        for index, value in enumerate(row_values):
            row[index].text = value
    set_table_geometry(table, widths)
    doc.add_paragraph()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def build_report() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Relatorio tecnico - VettiFlow + Protheus + Postgres")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = DARK_BLUE

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(
        f"Status local validado em {datetime.now().strftime('%d/%m/%Y %H:%M')}."
    )
    subtitle_run.font.color.rgb = MUTED
    subtitle_run.font.size = Pt(10)

    add_heading(doc, "Resumo executivo", 1)
    add_bullet(
        doc,
        "O VettiFlow consulta produtos reais do Protheus no Postgres local, usando SB1 para cadastro, SG1 para estrutura, SB2 para saldo e SC2 para OPs vinculadas.",
    )
    add_bullet(
        doc,
        "A criacao de OP no dashboard fica bloqueada quando o produto nao possui OP SMD vinculada na arvore Protheus ou quando algum componente material nao tem saldo suficiente.",
    )
    add_bullet(
        doc,
        "As telas de suporte e expedicao nao exibem mais OPs inventadas; elas dependem do fluxo real do ProductionFlowStore.",
    )
    add_bullet(
        doc,
        "Os relatorios administrativos foram testados no dashboard e renderizam os paineis de KPIs, tempos, defeitos, produto e detalhamento por OP.",
    )

    add_heading(doc, "Fluxo atual no Flutter", 1)
    add_table(
        doc,
        ["Etapa", "Como funciona hoje", "Fonte de dados"],
        [
            [
                "Autocomplete de produto",
                "O campo Codigo Protheus busca por codigo ou descricao e sugere produtos reais.",
                "protheus_raw.vw_sb1_products",
            ],
            [
                "Estrutura e ramificacoes",
                "Ao selecionar o produto, o app busca componentes diretos e OPs filhas; a liberacao SMD usa busca recursiva ate 4 niveis.",
                "SG1 + SC2",
            ],
            [
                "Validacao de estoque",
                "Para cada componente material, calcula quantidade necessaria = quantidade da OP x quantidade por unidade e compara com saldo disponivel.",
                "protheus_raw.vw_sb2_stock_balances",
            ],
            [
                "Movimentacao interna",
                "Criacao, avancos, pausas, fechamento, defeitos e eventos sao espelhados no schema vettiflow.",
                "vettiflow.production_*",
            ],
        ],
        [1.45, 3.2, 1.85],
    )

    add_heading(doc, "Regras de criacao de OP", 1)
    add_number(
        doc,
        "A Tatiane abre o dashboard e informa o codigo ou descricao do produto.",
    )
    add_number(
        doc,
        "O app consulta o Postgres e so habilita Criar OP se o produto existir na SB1.",
    )
    add_number(
        doc,
        "O app exige OP SMD vinculada na arvore do produto. Exemplo real: 730-0863 passa por 550-0863 e encontra 500-0863, que e SUB SMD SMART ALARM MONITORADA CENTRAL.",
    )
    add_number(
        doc,
        "O app valida estoque dos componentes materiais. Componentes de custo/mao de obra com codigo MOD... nao entram nessa checagem.",
    )
    add_number(
        doc,
        "Se faltar saldo, a OP nao e criada e a tela mostra o componente com quantidade necessaria e disponivel.",
    )

    add_heading(doc, "Validacao real observada", 1)
    add_table(
        doc,
        ["Item verificado", "Resultado"],
        [
            [
                "Produto 730-0863",
                "Encontrado na SB1 como SMART ALARM - MONITORADA CENTRAL.",
            ],
            [
                "Liberacao SMD",
                "Encontrada via ramificacao com OPs para 500-0863 / SUB SMD SMART ALARM MONITORADA CENTRAL.",
            ],
            [
                "Estoque",
                "A checagem local detectou bloqueio para 102-339: precisa 1 e o saldo estimado estava -2080 no momento do teste.",
            ],
            [
                "Relatorios admin",
                "Teste automatizado abriu Relatorios de Producao e verificou paineis de tempo por etapa, produto e detalhamento de OPs.",
            ],
            [
                "Dados fake",
                "OPs/produtos fake foram removidos de lib/. Fixtures ficaram somente em test/.",
            ],
        ],
        [2.1, 4.4],
    )

    add_heading(doc, "Como a FastAPI entraria", 1)
    paragraph = doc.add_paragraph()
    paragraph.add_run(
        "A FastAPI deveria ficar entre o Flutter e o banco/Protheus. "
        "Hoje o Flutter conecta direto no Postgres para acelerar o prototipo local. "
        "Na implantacao, o recomendado e o Flutter chamar uma API interna, e a API guardar credenciais, validar permissoes, registrar auditoria e consultar Postgres/Protheus."
    )

    add_table(
        doc,
        ["Camada", "Endereco exemplo", "Responsabilidade"],
        [
            [
                "Flutter VettiFlow",
                "http://192.168.10.80 ou app Windows na rede",
                "Interface de operadores, dashboard, criacao e movimentacao de OP.",
            ],
            [
                "FastAPI interna",
                "http://192.168.10.50:8000",
                "Endpoints /produtos, /ops, /estoque, /movimentos; validacao de usuario e regra de negocio.",
            ],
            [
                "Postgres no servidor",
                "192.168.10.50:5432/vettiflow",
                "Banco do VettiFlow e replica/exportacao reduzida das tabelas Protheus necessarias.",
            ],
            [
                "Protheus/SQL Server",
                "Servidor Protheus da empresa",
                "Fonte oficial para SB1, SC2, SG1, SB2 e futuras tabelas de apontamento/movimento.",
            ],
        ],
        [1.55, 1.95, 3.0],
    )

    add_heading(doc, "Credenciais e variaveis de ambiente", 1)
    add_table(
        doc,
        ["Variavel", "Exemplo", "Observacao"],
        [
            [
                "VETTIFLOW_API_BASE_URL",
                "http://192.168.10.50:8000",
                "URL que o Flutter usaria quando a FastAPI entrar.",
            ],
            [
                "VETTIFLOW_PG_HOST",
                "192.168.10.50",
                "IP do computador servidor do Protheus/Postgres.",
            ],
            ["VETTIFLOW_PG_PORT", "5432", "Porta padrao do Postgres."],
            ["VETTIFLOW_PG_DATABASE", "vettiflow", "Banco do VettiFlow."],
            [
                "VETTIFLOW_PG_USER",
                "vettiflow_app",
                "Usuario de aplicacao com permissoes limitadas.",
            ],
            [
                "VETTIFLOW_PG_PASSWORD",
                "********",
                "Nunca hardcoded no codigo; usar variavel de ambiente/secret.",
            ],
        ],
        [2.0, 2.1, 2.4],
    )

    add_heading(doc, "Endpoints FastAPI imaginados", 1)
    add_table(
        doc,
        ["Metodo", "Rota", "Uso"],
        [
            ["GET", "/produtos?busca=central", "Autocomplete por codigo/descricao."],
            ["GET", "/produtos/{codigo}", "Detalhe SB1 + estrutura SG1 + estoque SB2 + OPs SMD."],
            ["POST", "/ops", "Cria OP apos validar SMD, estoque e permissao da Tatiane."],
            ["PATCH", "/ops/{numero}/etapa", "Avanca, pausa, retoma ou conclui etapa."],
            ["GET", "/relatorios/producao", "KPIs e agregados usados pelo dashboard admin."],
        ],
        [1.0, 2.0, 3.5],
    )

    add_heading(doc, "Validacoes antes de colocar em producao", 1)
    add_bullet(doc, "Criar usuario Postgres exclusivo para o VettiFlow, sem usar postgres/superuser.")
    add_bullet(doc, "Abrir firewall apenas para a porta da FastAPI; idealmente o Flutter nao deve acessar 5432 diretamente.")
    add_bullet(doc, "Definir job de atualizacao das tabelas Protheus exportadas ou conexao controlada com a base oficial.")
    add_bullet(doc, "Confirmar com a gestao se o criterio oficial de liberacao SMD e codigo 500-/descricao SMD ou algum campo especifico do Protheus.")
    add_bullet(doc, "Criar backup automatico do banco vettiflow no servidor.")

    section = doc.add_section(WD_SECTION_START.CONTINUOUS)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("VettiFlow - relatorio tecnico de integracao")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_report()
    print(OUTPUT)
