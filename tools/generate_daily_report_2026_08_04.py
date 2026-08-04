from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "relatorio_vettiflow_mudancas_2026-08-04.docx"
LOGO = ROOT / "assets" / "images" / "vetti-flow-logo.png"

VETTI_BLUE = RGBColor(0, 119, 189)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(33, 43, 54)
MUTED = RGBColor(92, 107, 123)
GREEN = RGBColor(25, 135, 84)
ORANGE = RGBColor(190, 115, 25)
LIGHT_BLUE_FILL = "E8F4FB"
LIGHT_GRAY_FILL = "F2F4F7"
WARNING_FILL = "FFF4E0"
BORDER = "D8E0E8"


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, value: str = "nil", color: str = "FFFFFF") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), value)
        border.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if row_index == 0:
            tbl_header = tr_pr.find(qn("w:tblHeader"))
            if tbl_header is None:
                tr_pr.append(OxmlElement("w:tblHeader"))
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15


def set_paragraph_border_bottom(paragraph, color: str = "C8D6E2") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, VETTI_BLUE, 16, 8),
        ("Heading 2", 13, VETTI_BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header.paragraphs[0].paragraph_format.space_after = Pt(0)
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_borders(header_table)
    header_table.autofit = False
    left_cell = header_table.rows[0].cells[0]
    right_cell = header_table.rows[0].cells[1]
    left_cell.width = Inches(2.2)
    right_cell.width = Inches(4.3)
    set_cell_margins(left_cell, top=0, start=0, bottom=0, end=0)
    set_cell_margins(right_cell, top=0, start=0, bottom=0, end=0)
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    logo_p = left_cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO.exists():
        run = logo_p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.25))
    else:
        run = logo_p.add_run("VettiFlow")
        set_run_font(run, size=10, color=VETTI_BLUE, bold=True)

    top_right = right_cell.paragraphs[0]
    top_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    top_right.paragraph_format.space_after = Pt(8)
    run = top_right.add_run("VettiFlow | Relatorio interno")
    set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)

    subtitle = right_cell.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    subtitle.paragraph_format.space_after = Pt(0)
    run = subtitle.add_run("Atualizacoes do dia e movimentacoes Postgres/Protheus")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("VettiFlow 1.0 - relatorio operacional de progresso - 04/08/2026")
    set_run_font(run, size=9, color=MUTED)


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("RELATORIO DE MUDANCAS")
    set_run_font(run, size=10, color=MUTED, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("VettiFlow - atualizacoes do dia")
    set_run_font(run, size=24, color=DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run(
        "Resumo simples das melhorias feitas em 04/08/2026, com foco no fluxo de OP, assinaturas, privacidade e movimentacoes Protheus/Postgres."
    )
    set_run_font(run, size=12, color=MUTED)

    metadata = [
        ("Projeto", "VettiFlow 1.0 Flutter/Dart"),
        ("Data", "04/08/2026"),
        ("Objetivo", "Preparar o app para testes controlados com OPs reais e estoque espelhado do Protheus."),
        ("Validacao", "Analyzer limpo e testes automatizados principais passando."),
    ]
    info_table = doc.add_table(rows=len(metadata), cols=2)
    info_table.autofit = False
    set_table_borders(info_table, value="nil")
    for row_index, (label, value) in enumerate(metadata):
        label_cell = info_table.rows[row_index].cells[0]
        value_cell = info_table.rows[row_index].cells[1]
        label_cell.width = Inches(3.15)
        value_cell.width = Inches(3.35)
        set_cell_fill(label_cell, LIGHT_GRAY_FILL)
        set_cell_margins(label_cell, top=40, start=80, bottom=40, end=80)
        set_cell_margins(value_cell, top=40, start=80, bottom=40, end=80)
        label_p = label_cell.paragraphs[0]
        label_p.paragraph_format.space_after = Pt(0)
        label_run = label_p.add_run(label)
        set_run_font(label_run, size=10, color=INK, bold=True)
        value_p = value_cell.paragraphs[0]
        value_p.paragraph_format.space_after = Pt(0)
        value_run = value_p.add_run(value)
        set_run_font(value_run, size=10, color=INK)
    doc.add_paragraph()

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(6)
    rule.paragraph_format.space_after = Pt(10)
    set_paragraph_border_bottom(rule)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_fill(cell, fill)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    set_table_geometry(table, [6.5])
    doc.add_paragraph()


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for index, header in enumerate(headers):
        hdr[index].text = header
        set_cell_fill(hdr[index], LIGHT_GRAY_FILL)
        for paragraph in hdr[index].paragraphs:
            paragraph.runs[0].bold = True
            paragraph.runs[0].font.color.rgb = DARK_BLUE
    for row_values in rows:
        row = table.add_row().cells
        for index, value in enumerate(row_values):
            row[index].text = value
    set_table_geometry(table, widths)
    doc.add_paragraph()


def build_report() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    add_heading(doc, "Resumo executivo", 1)
    add_body(
        doc,
        "Hoje o VettiFlow foi ajustado para ficar mais confiavel na criacao, movimentacao e cancelamento de OPs. A ideia principal foi tirar a informalidade do papel e deixar o app registrar quem fez cada acao, em qual etapa, com qual PIN e quais tabelas do Protheus/Postgres foram afetadas.",
    )
    add_bullet(doc, "Criacao de OP agora exige assinatura quando gera movimento ligado ao Protheus.")
    add_bullet(doc, "Cancelamento de OP ficou mais claro e tambem exige PIN antes de devolver empenhos.")
    add_bullet(doc, "A OP ganhou uma area de auditoria com assinaturas e movimentacoes.")
    add_bullet(doc, "Gestores deixam de ver tempos, pausas e detalhes sensiveis de outros setores.")
    add_bullet(doc, "Relatorios passam a respeitar a area do gestor logado.")

    add_callout(
        doc,
        "Leitura para gestao",
        "O app esta deixando de ser apenas um painel visual e passando a funcionar como trilha operacional: ele registra origem, destino, responsavel e assinatura das acoes que impactam o fluxo e o estoque.",
    )

    add_heading(doc, "O que mudou para quem usa o app", 1)
    add_table(
        doc,
        ["Area", "Mudanca", "Efeito pratico"],
        [
            [
                "Criacao de OP",
                "A pessoa informa o produto Protheus, quantidade, armazem e PIN.",
                "Evita criacao anonima quando ha empenho/movimentacao ligada ao Protheus.",
            ],
            [
                "Escolha de armazem",
                "Quando falta item no armazem da OP, a tela mostra outros armazens com saldo e permite escolher um.",
                "A OP pode nascer usando material de outro armazem, mantendo a origem registrada.",
            ],
            [
                "SMD",
                "O SMD entrou antes da gravacao de firmware no fluxo.",
                "Paula/Leandro conseguem apontar SMD e liberar a OP para a producao.",
            ],
            [
                "Cancelamento",
                "A tela foi separada em um bloco mais legivel, com escolha de retorno por item e PIN obrigatorio.",
                "A devolucao de empenho fica rastreada e menos sujeita a erro manual.",
            ],
            [
                "Auditoria da OP",
                "O detalhe da OP mostra assinaturas de criacao, inicio, pausa/retomada e conclusao.",
                "Fica mais facil explicar quem fez o que e quando.",
            ],
            [
                "Privacidade",
                "Tempo, pausas e detalhes de materiais ficam ocultos para gestores de outros setores.",
                "Reduz exposicao desnecessaria e evita comparacoes/fofocas entre setores.",
            ],
        ],
        [1.35, 2.55, 2.6],
    )

    add_heading(doc, "Movimentacoes feitas no Protheus/Postgres", 1)
    add_body(
        doc,
        "Abaixo esta o comportamento atual do sistema. O Protheus aqui esta representado pelo banco Postgres que guarda as tabelas exportadas/espelhadas.",
    )
    add_table(
        doc,
        ["Momento", "O que o app grava/altera", "Observacao importante"],
        [
            [
                "Criar OP",
                "Cria empenhos na SD4, grava movimento RE0 na SD3 para itens fisicos e aumenta B2_QEMP na SB2.",
                "A estrutura vem da SG1. A SD4 nao e usada como fonte da estrutura porque pode estar desatualizada.",
            ],
            [
                "Criar OP com material de outro armazem",
                "O componente guarda o armazem escolhido. O empenho/movimento usa esse armazem de origem.",
                "Exemplo: OP no armazem 05 pode empenhar item que saiu do 01, se o usuario selecionar o 01.",
            ],
            [
                "Avancar etapa",
                "Grava movimento VFT na SD3. Se a etapa muda de armazem, atualiza B2_QATU na SB2: baixa origem e soma destino.",
                "Exemplo: Almoxarifado 01 para SMD 03 movimenta saldo atual do produto da OP entre armazens.",
            ],
            [
                "Produção para produção",
                "Grava SD3 VFT como auditoria, mas nao mexe no B2_QATU quando origem e destino sao o mesmo armazem.",
                "Gravacao, soldagem, teste e fechamento ficam no armazem 05.",
            ],
            [
                "Cancelar OP",
                "Cancela/baixa o empenho na SD4, grava auditoria de cancelamento, reduz B2_QEMP na SB2 e grava DE0 na SD3.",
                "O armazem de retorno escolhido fica registrado para auditoria. A acao atual devolve empenho; nao faz transferencia fisica extra de saldo atual.",
            ],
            [
                "MOD / Mao de obra",
                "Nao gera movimentacao fisica de estoque nem requisicao de armazem.",
                "MOD pode aparecer negativa no Protheus e e tratada como recurso nao limitante para estoque fisico.",
            ],
        ],
        [1.35, 3.1, 2.05],
    )

    add_heading(doc, "Mapa atual dos armazens", 1)
    add_table(
        doc,
        ["Armazem", "Setor", "Como entra no fluxo"],
        [
            ["01", "Almoxarifado", "Separacao inicial e origem de materiais da Vera."],
            ["03", "SMD", "Etapa da Paula/Leandro antes da gravacao de firmware."],
            ["05", "Producao", "Gravacao, soldagem, teste e fechamento."],
            ["10", "Expedicao", "Destino final de expedicao/armazenamento."],
            ["06/07", "Suporte", "Mapeado para o setor de suporte; ajustes mais finos ficam para a proxima etapa do fluxo."],
        ],
        [1.0, 1.6, 3.9],
    )

    add_heading(doc, "Regras de seguranca e permissao", 1)
    add_bullet(doc, "Tatiane e Andressa ficam focadas na producao; Tatiane tambem tem excecao para armazem 10.")
    add_bullet(doc, "Vera movimenta o almoxarifado e nao deve mexer em gravacao, teste ou outras etapas de producao.")
    add_bullet(doc, "Paula aponta SMD e libera para producao, mas nao deve criar OP no SMD.")
    add_bullet(doc, "Bruna e Tamara ficam no armazem 10/expedicao.")
    add_bullet(doc, "Bruno e Vinicius ficam ligados ao suporte.")
    add_bullet(doc, "Admin Master continua com permissao ampla para ajuste e acompanhamento geral.")

    add_heading(doc, "Regras de leitura do Protheus", 1)
    add_bullet(doc, "Produto bloqueado na SB1 pelo campo B1_MSBLQL = '1' nao aparece para operacao no app.")
    add_bullet(doc, "Produto disponivel na SB1 considera B1_MSBLQL = '2' ou sem bloqueio operacional.")
    add_bullet(doc, "A estrutura de componentes vem primeiro da SG1.")
    add_bullet(doc, "Saldo atual vem da SB2 usando B2_QATU. O app deixou de fazer conta de disponivel como saldo menos empenhado para exibir disponibilidade principal.")
    add_bullet(doc, "A SD4 foi removida como fallback de leitura de estrutura, pois pode estar desatualizada.")

    add_heading(doc, "Assinaturas e auditoria", 1)
    add_body(
        doc,
        "Foi adicionada uma visualizacao dentro da OP para mostrar as assinaturas das acoes feitas. O objetivo nao e expor pessoas, mas dar rastreabilidade operacional para o gestor entender o caminho da OP.",
    )
    add_table(
        doc,
        ["Acao", "Assinatura registrada", "Onde aparece"],
        [
            ["Criacao", "Nome do usuario e PIN mascarado.", "Detalhe da OP e payloads de movimento."],
            ["Inicio de etapa", "Operador, etapa e horario.", "Detalhe da OP."],
            ["Pausa/retomada", "Operador, motivo e horario.", "Detalhe da OP e relatorio do setor."],
            ["Conclusao/movimentacao", "Operador, PIN e etapa liberada.", "Detalhe da OP e SD3 VFT."],
            ["Cancelamento", "PIN obrigatorio e armazem escolhido por item.", "SD3/SD4/SB2 e detalhe de auditoria."],
        ],
        [1.5, 2.45, 2.55],
    )

    add_heading(doc, "O que ficou melhor para testar", 1)
    add_number(doc, "Criar uma OP no armazem 05 com componente vindo do 01 e conferir se SD4, SD3 e SB2 foram atualizadas.")
    add_number(doc, "Avancar uma OP do almoxarifado para SMD e conferir SD3 VFT e alteracao de B2_QATU de 01 para 03.")
    add_number(doc, "Avancar uma OP de SMD para gravacao e conferir SD3 VFT e alteracao de B2_QATU de 03 para 05.")
    add_number(doc, "Cancelar uma OP e conferir queda do B2_QEMP, cancelamento na SD4 e movimento DE0 na SD3.")
    add_number(doc, "Entrar com gestor de outro setor e confirmar que tempos/pausas/detalhes sensiveis nao aparecem.")

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_report()
    print(OUTPUT)
