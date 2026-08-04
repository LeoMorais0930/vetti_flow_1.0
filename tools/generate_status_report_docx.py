from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "relatorio_pratico_vettiflow_status_2026-07-31.docx"

VETTI_BLUE = RGBColor(0, 119, 189)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(33, 43, 54)
MUTED = RGBColor(92, 107, 123)
LIGHT_BLUE_FILL = "E8F4FB"
LIGHT_GRAY_FILL = "F2F4F7"
GREEN_FILL = "E7F6EC"
WARNING_FILL = "FFF4E0"
BORDER = "D8E0E8"


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15


def set_paragraph_border_bottom(paragraph, color: str = "C8D6E2") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, VETTI_BLUE, 16, 8),
        ("Heading 2", 13, VETTI_BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("VettiFlow 1.0 - relatorio pratico de progresso")
    set_run_font(run, size=9, color=MUTED)


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("RELATORIO DE PROJETO")
    set_run_font(run, size=10, color=MUTED, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("VettiFlow - o que foi feito ate agora")
    set_run_font(run, size=24, color=DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run(
        "Resumo pratico para acompanhamento, repasse e definicao dos proximos passos."
    )
    set_run_font(run, size=12, color=MUTED)

    metadata = [
        ("Projeto", "VettiFlow 1.0 Flutter/Dart"),
        ("Workspace", r"C:\Users\Leonardo Morais\Desktop\VettiFlow\vetti flow 1.0 flutter dart"),
        ("Data", "31/07/2026"),
        ("Validacao", "flutter test: 29 testes passaram"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, color=INK, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.5, color=INK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(6)
    rule.paragraph_format.space_after = Pt(10)
    set_paragraph_border_bottom(rule)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float],
    *,
    header_fill: str = LIGHT_GRAY_FILL,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_fill(header_cells[idx], header_fill)
        for paragraph in header_cells[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if paragraph.runs:
                set_run_font(paragraph.runs[0], size=10.5, color=INK, bold=True)

    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                if paragraph.runs:
                    set_run_font(paragraph.runs[0], size=10, color=INK)

    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_status_callout(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_fill(cell, GREEN_FILL)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(
        "Status geral: o prototipo ja possui fluxo operacional integrado entre dashboard, etapas de producao, "
        "armazenamento/expedicao, suporte, relatorios e uma primeira ponte local com dados Protheus em Postgres. "
        "A prioridade agora e trocar o acesso direto ao banco por uma API interna e validar com dados reais completos."
    )
    set_run_font(run, size=10.5, color=INK, bold=True)
    set_table_geometry(table, [6.5])
    doc.add_paragraph()


def build_report() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    add_heading(doc, "Resumo executivo", 1)
    add_status_callout(doc)
    add_bullet(doc, "Projeto Flutter/Dart fixado como workspace principal, com README e arquivo de referencia para evitar confusao com outras pastas.")
    add_bullet(doc, "App estruturado com Provider, RepositoryProvider e Cubit, mantendo uma fonte central de fluxo de producao.")
    add_bullet(doc, "Fluxo de OPs ampliado para cobrir criacao, etapas, pausas, assinatura por PIN, defeitos, armazenamento e expedicao.")
    add_bullet(doc, "Integracao local com Postgres/Protheus iniciada usando as tabelas SB1, SC2, SG1 e SB2 como base de produtos, OPs, estruturas e saldo.")
    add_bullet(doc, "Suite automatizada validada com 29 testes passando em 31/07/2026.")

    add_heading(doc, "Entregas concluidas", 1)
    add_table(
        doc,
        ["Frente", "O que foi feito", "Resultado pratico"],
        [
            [
                "Base do app",
                "Aplicacao Flutter organizada com tema Vetti, rotas, login, dashboard e telas por etapa.",
                "Projeto pronto para rodar em web, Windows e demais plataformas Flutter.",
            ],
            [
                "Fluxo de producao",
                "Criado modelo ProductionOrderFlow com etapas Almoxarifado, Gravacao, Soldagem, Teste, Fechamento, Expedicao, Armazenada e Finalizada.",
                "As telas passam a trabalhar com uma mesma OP ao longo do processo.",
            ],
            [
                "Pausas e tempos",
                "Adicionados status de execucao, sessoes por operador, eventos de pausa, motivos e duracao trabalhada.",
                "Gestao consegue rastrear tempo por etapa e por colaborador.",
            ],
            [
                "Dashboard",
                "Dashboard recebeu criacao de OP, visao de armazenadas, relatorios, detalhes e controles de avancar/voltar/cancelar.",
                "Gestao acompanha o fluxo sem depender apenas das telas operacionais.",
            ],
            [
                "Nova OP",
                "Dialogo de nova OP passou a buscar produto Protheus por codigo/descricao, preencher produto e mostrar componentes.",
                "Tatiane cria OP com dados mais proximos da base real.",
            ],
            [
                "Validacoes Protheus",
                "Criacao de OP e bloqueada sem OP SMD vinculada ou com componente material sem estoque suficiente.",
                "Evita iniciar producao sem liberacao ou sem saldo minimo.",
            ],
            [
                "Telas operacionais",
                "Almoxarifado, Expedicao e Suporte passaram a derivar filas do ProductionFlowStore em vez de dados soltos.",
                "Menos dados fake e mais consistencia entre telas.",
            ],
            [
                "Banco local",
                "Criadas migrations e views para schema vettiflow, metadados Protheus e area raw em jsonb.",
                "Postgres local preparado para prototipo e testes com export Protheus.",
            ],
        ],
        [1.45, 3.05, 2.0],
        header_fill=LIGHT_BLUE_FILL,
    )

    add_heading(doc, "Integracao Protheus e Postgres", 1)
    add_body(
        doc,
        "A integracao atual ainda e local, mas ja define a direcao tecnica: o app consulta produtos, estruturas, saldos "
        "e OPs Protheus a partir de views no Postgres. Isso permitiu testar regras reais antes de criar a API definitiva."
    )
    add_table(
        doc,
        ["Fonte", "Uso atual no VettiFlow"],
        [
            ["SB1", "Cadastro e busca de produtos por codigo ou descricao."],
            ["SG1", "Estrutura/componentes do produto para montar materiais da OP."],
            ["SB2", "Saldo disponivel estimado para bloquear criacao quando faltar componente."],
            ["SC2", "OPs vinculadas, incluindo liberacao SMD por ramificacao do produto."],
            ["vettiflow.production_*", "Persistencia propria do fluxo criado pelo VettiFlow: OPs, componentes, tempos, pausas, sessoes, defeitos e eventos."],
        ],
        [1.3, 5.2],
    )

    add_heading(doc, "Regras ja implementadas no fluxo", 1)
    add_number(doc, "Usuario informa codigo ou descricao do produto ao criar OP.")
    add_number(doc, "Sistema consulta o produto no Protheus/Postgres e mostra o produto selecionado.")
    add_number(doc, "Sistema carrega componentes e OPs filhas relacionadas.")
    add_number(doc, "Se for produto Protheus e nao existir OP SMD vinculada, a criacao e bloqueada.")
    add_number(doc, "Se algum componente material tiver saldo menor que o necessario, a criacao e bloqueada.")
    add_number(doc, "Componentes com codigo MOD... nao entram na validacao de estoque.")
    add_number(doc, "OP criada entra no fluxo do VettiFlow e passa pelas etapas operacionais.")

    add_heading(doc, "Telas e operacao", 1)
    add_table(
        doc,
        ["Tela", "Estado atual"],
        [
            ["Login", "Direciona operador para a tela correta conforme posto/atribuicao."],
            ["Dashboard/Gestao", "Cria OPs, acompanha status, responsaveis, armazenadas e relatorios."],
            ["Almoxarifado", "Mostra OPs na etapa de separacao, itens previstos e permite iniciar, pausar e entregar."],
            ["Firmware", "Fluxo operacional de gravacao com acoes, defeitos e assinatura."],
            ["Soldagem", "Tela operacional com inicio, pausa, conclusao e suporte a mais de um colaborador por OP."],
            ["Teste", "Registra defeitos e direciona demandas para suporte."],
            ["Fechamento", "Prepara OP para expedicao com quantidade fechada."],
            ["Expedicao", "Finaliza despacho, permite armazenar parcial ou total e expedir saldo armazenado depois."],
            ["Suporte Tecnico", "Lista OPs com defeitos vindos do teste e permite conferencia/requisicao."],
            ["TV", "Painel visual com rotacao de paineis de producao."],
        ],
        [1.75, 4.75],
    )

    add_heading(doc, "Banco, scripts e documentacao", 1)
    add_bullet(doc, "Criado diretorio db/ com migrations numeradas de 001 a 008.")
    add_bullet(doc, "Criadas views auxiliares para produtos SB1, estruturas SG1, saldos SB2, OPs SC2 e resumos de importacao.")
    add_bullet(doc, "Documentado o Postgres local, comandos para subir/parar servidor, aplicar migrations e importar export Protheus.")
    add_bullet(doc, "Criado manual rapido para colaboradores com passos de uso por tela.")
    add_bullet(doc, "Criadas ferramentas em tools/ para importacao, conferencia de sincronizacao e geracao de relatorio tecnico.")

    add_heading(doc, "Validacao realizada", 1)
    add_table(
        doc,
        ["Validacao", "Resultado"],
        [
            ["flutter test", "29 testes passaram em 31/07/2026."],
            ["Testes de Protheus", "Cobrem lookup de produto, componentes, OPs filhas, autocomplete, bloqueio sem SMD e bloqueio por estoque."],
            ["Testes de interface", "Cobrem rotas, login, telas principais, pausa, dashboard, relatorios, expedicao e TV."],
            ["Git status", "Ha mudancas locais ainda nao commitadas, incluindo arquivos novos em db/, docs/, tools/ e testes."],
        ],
        [1.8, 4.7],
        header_fill=GREEN_FILL,
    )

    add_heading(doc, "Pendencias e proximos passos", 1)
    add_table(
        doc,
        ["Prioridade", "Acao recomendada", "Por que importa"],
        [
            ["Alta", "Colocar uma FastAPI interna entre Flutter e Postgres/Protheus.", "Evita credenciais no app, centraliza regra de negocio e melhora seguranca."],
            ["Alta", "Criar usuario de banco limitado para a aplicacao.", "Nao usar postgres/superuser em ambiente real."],
            ["Alta", "Confirmar regra oficial de liberacao SMD no Protheus.", "Hoje a heuristica usa codigo 500-/descricao SMD; precisa validacao de negocio."],
            ["Media", "Fechar rotina de importacao/sincronizacao dos dados Protheus.", "Define se sera export, replica ou consulta controlada."],
            ["Media", "Validar fluxo completo com uma OP real de ponta a ponta.", "Garante aderencia da operacao real antes de implantar."],
            ["Media", "Revisar armazenamento de PIN.", "O codigo ja fala em hash no banco, mas a estrategia de seguranca precisa ser finalizada."],
            ["Baixa", "Limpar/organizar mudancas locais e fazer commit por frente.", "Facilita rastreio e continuidade do desenvolvimento."],
        ],
        [1.1, 3.0, 2.4],
        header_fill=WARNING_FILL,
    )

    add_heading(doc, "Arquivos principais alterados ou criados", 1)
    add_table(
        doc,
        ["Area", "Arquivos"],
        [
            ["App/configuracao", "lib/app/vetti_flow_app.dart, pubspec.yaml, README.md, PROJETO_FIXO_ATUAL.md"],
            ["Modelos/repositorios", "lib/data/models/protheus_product_lookup.dart, lib/data/repositories/protheus_product_repository.dart, production_flow_database.dart, flow_op_repository.dart, production_flow_store.dart"],
            ["Dashboard", "lib/ui/dashboard/dashboard_page.dart, cubit/dashboard_cubit.dart, widgets/nova_op_dialog.dart"],
            ["Operacao", "lib/ui/warehouse/warehouse_page.dart, lib/ui/expedition/expedition_page.dart, lib/ui/support/support_page.dart"],
            ["Banco", "db/migrations/*.sql, db/postgres/*.sql, db/README.md"],
            ["Testes", "test/protheus_lookup_test.dart, test/widget_test.dart, test/fakes/mock_op_repository.dart"],
            ["Documentacao/ferramentas", "docs/MANUAL_RAPIDO_COLABORADORES.md, tools/*.ps1, tools/*.dart, tools/*.py, tools/*.mjs"],
        ],
        [1.55, 4.95],
    )

    add_heading(doc, "Conclusao pratica", 1)
    add_body(
        doc,
        "O VettiFlow deixou de ser apenas uma interface visual e passou a ter um esqueleto operacional completo: "
        "fluxo de OP, telas por setor, regras de criacao, integracao inicial com dados Protheus, persistencia local em Postgres "
        "e testes automatizados. O proximo marco deve ser transformar a integracao local em arquitetura de producao com API, "
        "credenciais seguras e validacao com dados reais completos da Vetti."
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
